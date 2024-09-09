import sys
import requests
from bs4 import BeautifulSoup
from urllib.parse import urlparse
import os
import argparse
from concurrent.futures import ThreadPoolExecutor, as_completed
from docx import Document
from docx.shared import Inches
from io import BytesIO
import re
from unidecode import unidecode

def scrape_article(url):
    # Fetch the HTML content
    try:
        response = requests.get(url, timeout=10)
        response.raise_for_status()
    except requests.RequestException as e:
        return f"Error fetching {url}: {str(e)}"

    # Parse the HTML content
    soup = BeautifulSoup(response.text, 'html.parser')

    # Extract the title
    title_element = soup.find('h1')
    title = title_element.text.strip() if title_element else "Title not found"

    # Extract the content
    content_elements = soup.find_all('p')
    content = ' '.join([p.text.strip() for p in content_elements]) if content_elements else "Content not found"

    # Extract the updated date
    date_element = soup.find('time')
    updated_date = date_element['datetime'] if date_element and 'datetime' in date_element.attrs else 'Not found'

    # Extract the byline
    byline_element = soup.find(['span', 'div'], class_=lambda x: x and 'byline' in x.lower())
    byline = byline_element.text.strip() if byline_element else 'Not found'

    # Extract the section
    section_element = soup.find('meta', property='article:section')
    section = section_element['content'] if section_element else 'Not found'

    # Extract keywords/tags
    keywords = [meta['content'] for meta in soup.find_all('meta', property='article:tag')]

    # Extract image URL
    image_element = soup.find('meta', property='og:image')
    image_url = image_element['content'] if image_element else 'Not found'

    # Return the results
    return {
        'title': title,
        'content': content,
        'updated_date': updated_date,
        'byline': byline,
        'section': section,
        'keywords': keywords,
        'image_url': image_url,
        'url': url
    }

def sanitize_filename(title):
    # Remove non-ASCII characters
    title = unidecode(title)
    # Replace spaces with underscores and remove any other non-word characters
    return re.sub(r'[^\w\s-]', '', title.replace(' ', '_')).strip('-_')

def save_to_docx(data, filename):
    doc = Document()
    doc.add_heading(data['title'], 0)

    doc.add_paragraph(f"By: {data['byline']}")
    doc.add_paragraph(f"Updated: {data['updated_date']}")
    doc.add_paragraph(f"Section: {data['section']}")
    doc.add_paragraph(f"Keywords: {', '.join(data['keywords'])}")

    if data['image_url'] != 'Not found':
        try:
            response = requests.get(data['image_url'])
            image = BytesIO(response.content)
            doc.add_picture(image, width=Inches(6))
        except Exception as e:
            doc.add_paragraph(f"Error adding image: {str(e)}")

    doc.add_paragraph(data['content'])

    doc.add_paragraph(f"Source: {data['url']}")

    os.makedirs(os.path.dirname(filename), exist_ok=True)
    doc.save(filename)

def scrape_multiple_urls(urls, output_dir):
    results = []
    with ThreadPoolExecutor(max_workers=5) as executor:
        future_to_url = {executor.submit(scrape_article, url): url for url in urls}
        for future in as_completed(future_to_url):
            url = future_to_url[future]
            try:
                data = future.result()
                if isinstance(data, str):  # Error occurred
                    print(data)
                else:
                    results.append(data)
                    # Use the article title for the filename
                    safe_title = sanitize_filename(data['title'])
                    filename = os.path.join(output_dir, f"{safe_title[:100]}.docx")
                    save_to_docx(data, filename)
                    print(f"Scraped and saved: {url}")
            except Exception as exc:
                print(f"{url} generated an exception: {exc}")
    return results

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Scrape articles from URLs and save as Word documents")
    parser.add_argument("urls", nargs="+", help="URLs to scrape")
    parser.add_argument("-o", "--output", default="scraped_articles", help="Output directory for scraped data")
    args = parser.parse_args()

    results = scrape_multiple_urls(args.urls, args.output)
    print(f"\nScraped {len(results)} articles successfully.")
